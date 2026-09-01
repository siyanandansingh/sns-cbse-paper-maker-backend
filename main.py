import os, io, json, uuid, datetime, re
from pathlib import Path
from typing import Optional

import requests
import qrcode
from fastapi import FastAPI, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from sqlalchemy import create_engine, Column, Integer, String, Boolean, DateTime, Text, ForeignKey
from sqlalchemy.orm import declarative_base, sessionmaker, Session
from passlib.context import CryptContext
from jose import jwt
from dotenv import load_dotenv
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm

load_dotenv()
BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "data" / "app.db"
DB_PATH.parent.mkdir(exist_ok=True)

SECRET_KEY = os.getenv("SECRET_KEY", "change-me")
ADMIN_ID = os.getenv("ADMIN_ID", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "CHANGE_ME")
UPI_ID = os.getenv("UPI_ID", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")
DATABASE_URL = os.getenv("DATABASE_URL", "").strip()

if DATABASE_URL:
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = "postgresql+psycopg://" + DATABASE_URL[len("postgres://"):]
    elif DATABASE_URL.startswith("postgresql://"):
        DATABASE_URL = "postgresql+psycopg://" + DATABASE_URL[len("postgresql://"):]
    engine = create_engine(DATABASE_URL, pool_pre_ping=True)
else:
    engine = create_engine(f"sqlite:///{DB_PATH}", connect_args={"check_same_thread": False})

SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()
pwd = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")
bearer = HTTPBearer(auto_error=False)

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    mobile = Column(String, unique=True, index=True)
    name = Column(String, default="")
    password_hash = Column(String)
    is_admin = Column(Boolean, default=False)
    is_active = Column(Boolean, default=True)
    credits = Column(Integer, default=5)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class Payment(Base):
    __tablename__ = "payments"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    order_ref = Column(String, unique=True, index=True)
    amount = Column(Integer)
    credits = Column(Integer)
    utr = Column(String, default="")
    status = Column(String, default="PENDING")
    created_at = Column(DateTime, default=datetime.datetime.utcnow)
    decided_at = Column(DateTime, nullable=True)

class Paper(Base):
    __tablename__ = "papers"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, ForeignKey("users.id"))
    title = Column(String)
    class_name = Column(String)
    subject = Column(String)
    subject_code = Column(String)
    exam_type = Column(String)
    max_marks = Column(Integer)
    syllabus_source = Column(String)
    syllabus_text = Column(Text)
    question_paper = Column(Text)
    answer_key = Column(Text)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

class Setting(Base):
    __tablename__ = "settings"
    key = Column(String, primary_key=True)
    value = Column(Text)

Base.metadata.create_all(engine)

def db():
    s = SessionLocal()
    try:
        yield s
    finally:
        s.close()

def get_setting(s: Session, key: str, default: str = "") -> str:
    row = s.query(Setting).filter(Setting.key == key).first()
    return row.value if row else default

def put_setting(s: Session, key: str, value: str):
    row = s.query(Setting).filter(Setting.key == key).first()
    if row:
        row.value = value
    else:
        s.add(Setting(key=key, value=value))

def seed():
    s = SessionLocal()
    try:
        a = s.query(User).filter(User.mobile == ADMIN_ID).first()
        if not a:
            s.add(User(mobile=ADMIN_ID, name="ADMIN", password_hash=pwd.hash(ADMIN_PASSWORD), is_admin=True, credits=999999999))
        for k, v in {
            "free_credits": "5",
            "plan_10_credits": "3",
            "plan_20_credits": "10",
            "upi_id": UPI_ID,
            "app_name": "SNS CBSE Paper Maker",
        }.items():
            if not s.query(Setting).filter(Setting.key == k).first():
                s.add(Setting(key=k, value=v))
        s.commit()
    finally:
        s.close()
seed()

app = FastAPI(title="SNS CBSE Paper Maker API", version="2.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=False, allow_methods=["*"], allow_headers=["*"])

class LoginIn(BaseModel):
    mobile: str
    password: str
class RegisterIn(BaseModel):
    mobile: str
    name: str = ""
    password: str
class PaperIn(BaseModel):
    class_name: str
    subject: str
    subject_code: str = ""
    exam_type: str
    max_marks: int = 80
    difficulty: str = "Moderate"
    language: str = "English"
    syllabus_text: str = ""
    syllabus_source: str = "CBSE_DEFAULT"
class PaymentCreate(BaseModel):
    amount: int
class PaymentSubmit(BaseModel):
    order_ref: str
    utr: str
class CreditChange(BaseModel):
    amount: int
class AdminSettingsIn(BaseModel):
    free_credits: int
    plan_10_credits: int
    plan_20_credits: int
    upi_id: str

def token_for(u: User):
    return jwt.encode({"uid": u.id, "admin": u.is_admin, "exp": datetime.datetime.utcnow() + datetime.timedelta(days=30)}, SECRET_KEY, algorithm="HS256")

def current_user(c: Optional[HTTPAuthorizationCredentials] = Depends(bearer), s: Session = Depends(db)):
    if not c:
        raise HTTPException(401, "Login required")
    try:
        p = jwt.decode(c.credentials, SECRET_KEY, algorithms=["HS256"])
        u = s.get(User, int(p["uid"]))
    except Exception:
        u = None
    if not u or not u.is_active:
        raise HTTPException(401, "Invalid account")
    return u

def admin_user(u: User = Depends(current_user)):
    if not u.is_admin:
        raise HTTPException(403, "Admin only")
    return u

@app.get("/")
def root():
    return {"name": "SNS CBSE Paper Maker", "status": "online"}

@app.get("/health")
def health(s: Session = Depends(db)):
    return {"ok": True, "database": "ok", "ai_configured": bool(GEMINI_API_KEY), "users": s.query(User).count()}

@app.post("/auth/register")
def register(x: RegisterIn, s: Session = Depends(db)):
    mobile = re.sub(r"\s+", "", x.mobile)
    if len(mobile) < 6 or len(x.password) < 4:
        raise HTTPException(400, "Enter valid mobile/ID and password")
    if s.query(User).filter(User.mobile == mobile).first():
        raise HTTPException(400, "Mobile already registered")
    free = int(get_setting(s, "free_credits", "5"))
    u = User(mobile=mobile, name=x.name.strip() or "Teacher", password_hash=pwd.hash(x.password), credits=free)
    s.add(u); s.commit(); s.refresh(u)
    return {"token": token_for(u), "is_admin": False, "credits": u.credits, "name": u.name}

@app.post("/auth/login")
def login(x: LoginIn, s: Session = Depends(db)):
    u = s.query(User).filter(User.mobile == x.mobile.strip()).first()
    if not u or not pwd.verify(x.password, u.password_hash):
        raise HTTPException(401, "Wrong ID or password")
    if not u.is_active:
        raise HTTPException(403, "Account disabled by admin")
    return {"token": token_for(u), "is_admin": u.is_admin, "credits": "UNLIMITED" if u.is_admin else u.credits, "name": u.name}

@app.get("/me")
def me(u: User = Depends(current_user)):
    return {"mobile": u.mobile, "name": u.name, "is_admin": u.is_admin, "credits": "UNLIMITED" if u.is_admin else u.credits}

def syllabus_context(x: PaperIn, s: Session):
    key = f"syllabus::{x.class_name}::{x.subject_code or x.subject}".lower()
    custom = get_setting(s, key, "").strip()
    if x.syllabus_text.strip():
        return x.syllabus_text.strip(), False
    if custom:
        return custom, False
    return (f"Use the current official CBSE curriculum/syllabus for Class {x.class_name}, Subject {x.subject}, "
            f"Subject Code {x.subject_code}. Prefer official CBSE Academic sources. Do not include out-of-syllabus content."), True

def fallback_paper(x: PaperIn, syllabus: str):
    qp = f"""SNS CBSE PAPER MAKER\nCBSE STYLE QUESTION PAPER\nClass: {x.class_name}\nSubject: {x.subject} ({x.subject_code})\nExamination: {x.exam_type}\nMaximum Marks: {x.max_marks}\n\nAI generation is temporarily unavailable. Please try again when the free AI service is configured/available.\n\nSyllabus reference:\n{syllabus[:2500]}"""
    ak = "ANSWER KEY / MARKING SCHEME\nAI generation is temporarily unavailable. Please regenerate the paper."
    return qp, ak

def gemini_generate(x: PaperIn, syllabus: str, use_search: bool):
    if not GEMINI_API_KEY:
        return fallback_paper(x, syllabus)
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    prompt = f"""Create a complete, original CBSE-style school question paper AND answer key/marking scheme.
Class: {x.class_name}
Subject: {x.subject}
Subject Code: {x.subject_code}
Examination: {x.exam_type}
Maximum Marks: {x.max_marks}
Difficulty: {x.difficulty}
Language: {x.language}
Syllabus/source: {syllabus}

Rules:
1. If official current CBSE syllabus must be determined, use only reliable official CBSE Academic information and stay strictly in syllabus.
2. Follow the current CBSE-style competency/application focus appropriate for the class/subject.
3. Include objective, short/long answer, competency/case/source-based questions where appropriate.
4. Use sensible internal choices where appropriate.
5. The grand total MUST equal exactly {x.max_marks} marks.
6. Create original questions; do not copy long textbook passages verbatim.
7. Answer key must map to every question and give step-wise marking points.
8. Return ONLY valid JSON with keys question_paper and answer_key."""
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"responseMimeType": "application/json", "temperature": 0.45}
    }
    if use_search:
        payload["tools"] = [{"google_search": {}}]
    try:
        r = requests.post(url, json=payload, timeout=120)
        r.raise_for_status()
        data = r.json()
        text = "".join(p.get("text", "") for p in data["candidates"][0]["content"]["parts"] if "text" in p).strip()
        text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text)
        obj = json.loads(text)
        return obj["question_paper"], obj["answer_key"]
    except Exception as e:
        print("GEMINI ERROR:", repr(e), flush=True)
        return fallback_paper(x, syllabus)

@app.post("/papers/generate")
def generate(x: PaperIn, u: User = Depends(current_user), s: Session = Depends(db)):
    if x.max_marks < 10 or x.max_marks > 200:
        raise HTTPException(400, "Maximum marks must be between 10 and 200")
    if not u.is_admin and u.credits <= 0:
        raise HTTPException(402, "No paper credits. Buy Ã¢â€šÂ¹10 or Ã¢â€šÂ¹20 plan.")
    syllabus, use_search = syllabus_context(x, s)
    qp, ak = gemini_generate(x, syllabus, use_search)
    title = f"Class {x.class_name} {x.subject} {x.exam_type}"
    p = Paper(user_id=u.id, title=title, class_name=x.class_name, subject=x.subject, subject_code=x.subject_code,
              exam_type=x.exam_type, max_marks=x.max_marks, syllabus_source=x.syllabus_source if x.syllabus_text else "CBSE_DEFAULT",
              syllabus_text=syllabus, question_paper=qp, answer_key=ak)
    s.add(p)
    if not u.is_admin:
        u.credits -= 1
    s.commit(); s.refresh(p)
    return {"paper_id": p.id, "question_paper": qp, "answer_key": ak, "credits": "UNLIMITED" if u.is_admin else u.credits}

@app.get("/papers")
def papers(u: User = Depends(current_user), s: Session = Depends(db)):
    rows = s.query(Paper).filter(Paper.user_id == u.id).order_by(Paper.id.desc()).all()
    return [{"id": p.id, "title": p.title, "created_at": p.created_at.isoformat()} for p in rows]

@app.get("/papers/{pid}")
def paper(pid: int, u: User = Depends(current_user), s: Session = Depends(db)):
    p = s.get(Paper, pid)
    if not p or (p.user_id != u.id and not u.is_admin):
        raise HTTPException(404, "Paper not found")
    return {"id": p.id, "title": p.title, "question_paper": p.question_paper, "answer_key": p.answer_key}

@app.post("/payments/create")
def payment_create(x: PaymentCreate, u: User = Depends(current_user), s: Session = Depends(db)):
    if u.is_admin:
        raise HTTPException(400, "Admin does not need payment")
    if x.amount not in (10, 20):
        raise HTTPException(400, "Invalid plan")
    credits = int(get_setting(s, "plan_10_credits" if x.amount == 10 else "plan_20_credits", "3" if x.amount == 10 else "10"))
    upi = get_setting(s, "upi_id", UPI_ID).strip()
    if not upi:
        raise HTTPException(503, "UPI ID not configured by admin")
    ref = "SNS" + uuid.uuid4().hex[:12].upper()
    p = Payment(user_id=u.id, order_ref=ref, amount=x.amount, credits=credits)
    s.add(p); s.commit()
    uri = f"upi://pay?pa={upi}&pn=SNS%20CBSE%20Paper%20Maker&am={x.amount}.00&cu=INR&tn={ref}&tr={ref}"
    img = qrcode.make(uri); b = io.BytesIO(); img.save(b, format="PNG")
    import base64
    return {"order_ref": ref, "amount": x.amount, "credits": credits, "upi_uri": uri, "qr_base64": base64.b64encode(b.getvalue()).decode()}

@app.post("/payments/submit")
def payment_submit(x: PaymentSubmit, u: User = Depends(current_user), s: Session = Depends(db)):
    p = s.query(Payment).filter(Payment.order_ref == x.order_ref, Payment.user_id == u.id).first()
    if not p:
        raise HTTPException(404, "Payment order not found")
    if p.status != "PENDING":
        raise HTTPException(400, "Payment already processed")
    utr = x.utr.strip()
    if len(utr) < 6:
        raise HTTPException(400, "Enter valid UTR/transaction ID")
    other = s.query(Payment).filter(Payment.utr == utr, Payment.status.in_(["PENDING", "APPROVED"]), Payment.id != p.id).first()
    if other:
        raise HTTPException(400, "This UTR is already submitted")
    p.utr = utr; s.commit()
    return {"status": "PENDING", "message": "Payment sent for admin verification."}

@app.get("/payments/my")
def my_payments(u: User = Depends(current_user), s: Session = Depends(db)):
    rows = s.query(Payment).filter(Payment.user_id == u.id).order_by(Payment.id.desc()).all()
    return [{"id": p.id, "order_ref": p.order_ref, "amount": p.amount, "credits": p.credits, "utr": p.utr,
             "status": p.status, "created_at": p.created_at.isoformat()} for p in rows]

def make_docx(p: Paper):
    d = Document(); d.add_heading(p.title, 0); d.add_paragraph(p.question_paper); d.add_page_break(); d.add_heading("ANSWER KEY / MARKING SCHEME", 1); d.add_paragraph(p.answer_key)
    b = io.BytesIO(); d.save(b); b.seek(0); return b

def make_pdf(p: Paper):
    b = io.BytesIO(); c = canvas.Canvas(b, pagesize=A4); W, H = A4; x = 18*mm; y = H-18*mm
    def put(text):
        nonlocal y
        for raw in text.splitlines():
            line = raw or " "
            while len(line) > 95:
                c.drawString(x, y, line[:95]); line = line[95:]; y -= 5*mm
                if y < 18*mm: c.showPage(); c.setFont("Helvetica", 9); y = H-18*mm
            c.drawString(x, y, line); y -= 5*mm
            if y < 18*mm: c.showPage(); c.setFont("Helvetica", 9); y = H-18*mm
    c.setFont("Helvetica-Bold", 14); c.drawString(x, y, p.title); y -= 8*mm; c.setFont("Helvetica", 9); put(p.question_paper)
    c.showPage(); y = H-18*mm; c.setFont("Helvetica-Bold", 13); c.drawString(x, y, "ANSWER KEY / MARKING SCHEME"); y -= 8*mm; c.setFont("Helvetica", 9); put(p.answer_key); c.save(); b.seek(0); return b

@app.get("/papers/{pid}/docx")
def download_docx(pid: int, u: User = Depends(current_user), s: Session = Depends(db)):
    p = s.get(Paper, pid)
    if not p or (p.user_id != u.id and not u.is_admin): raise HTTPException(404)
    return StreamingResponse(make_docx(p), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="SNS_paper_{pid}.docx"'})

@app.get("/papers/{pid}/pdf")
def download_pdf(pid: int, u: User = Depends(current_user), s: Session = Depends(db)):
    p = s.get(Paper, pid)
    if not p or (p.user_id != u.id and not u.is_admin): raise HTTPException(404)
    return StreamingResponse(make_pdf(p), media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="SNS_paper_{pid}.pdf"'})

# ---------------- IN-APP ADMIN API ----------------
@app.get("/admin/api/stats")
def admin_stats(a: User = Depends(admin_user), s: Session = Depends(db)):
    return {
        "users": s.query(User).filter(User.is_admin == False).count(),
        "papers": s.query(Paper).count(),
        "pending_payments": s.query(Payment).filter(Payment.status == "PENDING", Payment.utr != "").count(),
        "approved_revenue": sum(p.amount for p in s.query(Payment).filter(Payment.status == "APPROVED").all()),
    }

@app.get("/admin/api/users")
def admin_users(a: User = Depends(admin_user), s: Session = Depends(db)):
    rows = s.query(User).filter(User.is_admin == False).order_by(User.id.desc()).all()
    return [{"id": u.id, "mobile": u.mobile, "name": u.name, "credits": u.credits, "is_active": u.is_active, "created_at": u.created_at.isoformat()} for u in rows]

@app.get("/admin/api/payments")
def admin_payments(a: User = Depends(admin_user), s: Session = Depends(db)):
    rows = s.query(Payment).order_by(Payment.id.desc()).limit(200).all()
    out = []
    for p in rows:
        u = s.get(User, p.user_id)
        out.append({"id": p.id, "teacher": u.name if u else "", "mobile": u.mobile if u else "", "amount": p.amount, "credits": p.credits,
                    "utr": p.utr, "status": p.status, "created_at": p.created_at.isoformat()})
    return out

@app.post("/admin/api/payment/{pid}/{decision}")
def admin_decide_payment(pid: int, decision: str, a: User = Depends(admin_user), s: Session = Depends(db)):
    p = s.get(Payment, pid)
    if not p or p.status != "PENDING": raise HTTPException(400, "Payment already processed/not found")
    if decision == "approve":
        if not p.utr: raise HTTPException(400, "UTR is missing")
        u = s.get(User, p.user_id); u.credits += p.credits; p.status = "APPROVED"
    elif decision == "reject": p.status = "REJECTED"
    else: raise HTTPException(400, "Invalid decision")
    p.decided_at = datetime.datetime.utcnow(); s.commit()
    return {"status": p.status}

@app.post("/admin/api/user/{uid}/credits")
def admin_change_credits(uid: int, x: CreditChange, a: User = Depends(admin_user), s: Session = Depends(db)):
    u = s.get(User, uid)
    if not u or u.is_admin: raise HTTPException(404)
    u.credits = max(0, u.credits + x.amount); s.commit(); return {"credits": u.credits}

@app.post("/admin/api/user/{uid}/toggle")
def admin_toggle_user(uid: int, a: User = Depends(admin_user), s: Session = Depends(db)):
    u = s.get(User, uid)
    if not u or u.is_admin: raise HTTPException(404)
    u.is_active = not u.is_active; s.commit(); return {"is_active": u.is_active}

@app.get("/admin/api/settings")
def admin_get_settings(a: User = Depends(admin_user), s: Session = Depends(db)):
    return {"free_credits": int(get_setting(s, "free_credits", "5")), "plan_10_credits": int(get_setting(s, "plan_10_credits", "3")),
            "plan_20_credits": int(get_setting(s, "plan_20_credits", "10")), "upi_id": get_setting(s, "upi_id", UPI_ID)}

@app.post("/admin/api/settings")
def admin_save_settings(x: AdminSettingsIn, a: User = Depends(admin_user), s: Session = Depends(db)):
    if min(x.free_credits, x.plan_10_credits, x.plan_20_credits) < 0: raise HTTPException(400, "Credits cannot be negative")
    put_setting(s, "free_credits", str(x.free_credits)); put_setting(s, "plan_10_credits", str(x.plan_10_credits)); put_setting(s, "plan_20_credits", str(x.plan_20_credits)); put_setting(s, "upi_id", x.upi_id.strip()); s.commit()
    return {"status": "saved"}
